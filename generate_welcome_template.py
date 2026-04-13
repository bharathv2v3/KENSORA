"""
Kensora by Kalakruthi — Client Welcome & Project Proposal Document Generator
Generates a professional .docx template with:
  - Welcome section
  - Project details & requirements
  - Room-by-room 3D & AR model placeholders with use-case descriptions
  - Kensora design recommendations per room
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Brand Colors ──────────────────────────────────────────────────────────────
TEAL_DARK   = RGBColor(0x0D, 0x3D, 0x3A)   # #0d3d3a
TEAL        = RGBColor(0x0F, 0x4F, 0x4B)   # #0f4f4b
GOLD        = RGBColor(0xB5, 0x9A, 0x6A)   # #b59a6a
GOLD_LIGHT  = RGBColor(0xD4, 0xB8, 0x8C)   # #d4b88c
CREAM       = RGBColor(0xF5, 0xF0, 0xE8)   # #f5f0e8
CHARCOAL    = RGBColor(0x1A, 0x1A, 0x1A)   # #1a1a1a
TEXT_MID    = RGBColor(0x5A, 0x55, 0x4E)   # #5a554e
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# ─── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set solid background color for a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def add_border(table, color_hex='B59A6A', size='6'):
    """Apply outer borders to a table."""
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'),   'single')
        b.set(qn('w:sz'),    size)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color_hex)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        m = OxmlElement(f'w:{side}')
        m.set(qn('w:w'),    str(val))
        m.set(qn('w:type'), 'dxa')
        tcMar.append(m)
    tcPr.append(tcMar)


def heading_run(para, text, size_pt, color, bold=False, italic=False, spacing_after=80):
    run = para.add_run(text)
    run.font.size   = Pt(size_pt)
    run.font.color.rgb = color
    run.font.bold   = bold
    run.font.italic = italic
    run.font.name   = 'Palatino Linotype'
    para.paragraph_format.space_after = Pt(0)
    return run


def add_gold_rule(doc, width_pct=80):
    """Add a thin gold horizontal rule as a single-cell table."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'B59A6A')
    cell.paragraphs[0].paragraph_format.space_before = Pt(0)
    cell.paragraphs[0].paragraph_format.space_after  = Pt(0)
    # Height via trPr
    tr  = tbl.rows[0]._tr
    trPr = OxmlElement('w:trPr')
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '18')   # ~1 px
    trPr.append(trH)
    tr.insert(0, trPr)
    # Width
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    '7000')
    tblW.set(qn('w:type'), 'dxa')
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    tblPr.append(tblW)
    doc.add_paragraph()   # spacing below rule


def body_para(doc, text, size_pt=10.5, color=None, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, space_after=6, space_before=0):
    if color is None:
        color = TEXT_MID
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text)
    r.font.size  = Pt(size_pt)
    r.font.color.rgb = color
    r.font.bold  = bold
    r.font.italic = italic
    r.font.name  = 'Calibri'
    return p


def label_para(doc, text, size_pt=8, color=None):
    """Small uppercase gold label."""
    if color is None:
        color = GOLD
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text.upper())
    r.font.size        = Pt(size_pt)
    r.font.color.rgb   = color
    r.font.bold        = True
    r.font.name        = 'Calibri'
    # letter spacing via rPr spacing
    rPr = r._r.get_or_add_rPr()
    sp  = OxmlElement('w:spacing')
    sp.set(qn('w:val'), '30')
    rPr.append(sp)
    return p


def add_placeholder_box(doc, label, description):
    """
    Creates a teal-bordered placeholder box for 3D / AR model.
    label       — e.g. '3D Model Placeholder'
    description — short use-case text
    """
    tbl = doc.add_table(rows=3, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_border(tbl, color_hex='0F4F4B', size='8')

    # Row 0 — label header (teal bg)
    header_cell = tbl.cell(0, 0)
    set_cell_bg(header_cell, '0F4F4B')
    set_cell_margins(header_cell, top=100, bottom=100, left=160, right=160)
    hp = header_cell.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run(f'  ◈  {label.upper()}  ◈  ')
    hr.font.size      = Pt(9)
    hr.font.color.rgb = GOLD_LIGHT
    hr.font.bold      = True
    hr.font.name      = 'Calibri'

    # Row 1 — insert image / blank box instruction
    img_cell = tbl.cell(1, 0)
    set_cell_bg(img_cell, 'F0EDE7')
    set_cell_margins(img_cell, top=200, bottom=200, left=200, right=200)
    # Set row height
    tr  = tbl.rows[1]._tr
    trPr = OxmlElement('w:trPr')
    trH  = OxmlElement('w:trHeight')
    trH.set(qn('w:val'), '1800')    # ~3.2 cm
    trPr.append(trH)
    tr.insert(0, trPr)
    ip = img_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.paragraph_format.space_before = Pt(12)
    ir = ip.add_run('[ INSERT RENDERED IMAGE / SCREENSHOT / AR QR CODE HERE ]')
    ir.font.size      = Pt(9)
    ir.font.color.rgb = RGBColor(0xBB, 0xB5, 0xAC)
    ir.font.italic    = True
    ir.font.name      = 'Calibri'

    # Row 2 — description
    desc_cell = tbl.cell(2, 0)
    set_cell_bg(desc_cell, 'FDFBF8')
    set_cell_margins(desc_cell, top=120, bottom=120, left=160, right=160)
    dp = desc_cell.paragraphs[0]
    dp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    dr = dp.add_run(description)
    dr.font.size      = Pt(9.5)
    dr.font.color.rgb = TEXT_MID
    dr.font.italic    = True
    dr.font.name      = 'Calibri'

    doc.add_paragraph()   # breathing room


def add_room_section(doc, room_name, room_icon,
                     requirements, usecase_3d, usecase_ar,
                     recommendations):
    """Full room block: heading + 2 placeholders + recommendations."""

    # ── Room Heading ──────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(f'{room_icon}  {room_name}')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    # ── Requirements ──────────────────────────────────────────────────────
    label_para(doc, 'Client Requirements')
    for req in requirements:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(req)
        r.font.size      = Pt(10)
        r.font.color.rgb = TEXT_MID
        r.font.name      = 'Calibri'

    # ── 3D Model Placeholder ──────────────────────────────────────────────
    label_para(doc, '3D Model Visualization')
    add_placeholder_box(doc,
        label='3D Model Placeholder',
        description=usecase_3d)

    # ── AR Model Placeholder ──────────────────────────────────────────────
    label_para(doc, 'Augmented Reality (AR) Model')
    add_placeholder_box(doc,
        label='AR Model Placeholder',
        description=usecase_ar)

    # ── Kensora Recommendations ───────────────────────────────────────────
    label_para(doc, 'Kensora Design Recommendations')
    for rec in recommendations:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after  = Pt(3)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(rec)
        r.font.size      = Pt(10)
        r.font.color.rgb = TEXT_MID
        r.font.name      = 'Calibri'

    doc.add_paragraph()   # section gap


# ════════════════════════════════════════════════════════════════════════════════
#  MAIN DOCUMENT BUILD
# ════════════════════════════════════════════════════════════════════════════════

def build_document():
    doc = Document()

    # ── Page Margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width    = Cm(21)
        section.page_height   = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)

    # ════════════════════════════════════════════════════════════════════════
    #  LETTERHEAD HEADER (2-col table: logo text | company info)
    # ════════════════════════════════════════════════════════════════════════
    tbl = doc.add_table(rows=1, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left cell — Brand
    lc = tbl.cell(0, 0)
    set_cell_bg(lc, '0D3D3A')
    set_cell_margins(lc, top=220, bottom=200, left=300, right=80)
    lp1 = lc.paragraphs[0]
    lp1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr1 = lp1.add_run('KENSORA')
    lr1.font.size      = Pt(28)
    lr1.font.color.rgb = WHITE
    lr1.font.bold      = True
    lr1.font.name      = 'Palatino Linotype'

    lp2 = lc.add_paragraph()
    lp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr2 = lp2.add_run('by Kalakruthi')
    lr2.font.size      = Pt(10)
    lr2.font.color.rgb = GOLD_LIGHT
    lr2.font.italic    = True
    lr2.font.name      = 'Calibri'

    lp3 = lc.add_paragraph()
    lp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    lr3 = lp3.add_run('◈  ULTRA-PREMIUM INTERIORS')
    lr3.font.size      = Pt(7.5)
    lr3.font.color.rgb = RGBColor(0xB5, 0x9A, 0x6A)
    lr3.font.bold      = True
    lr3.font.name      = 'Calibri'

    # Right cell — Company info
    rc = tbl.cell(0, 1)
    set_cell_bg(rc, '0D3D3A')
    set_cell_margins(rc, top=220, bottom=200, left=80, right=300)
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr = rp.add_run(
        'Kensora by Kalakruthi\n'
        'Design Studio, Indiranagar\n'
        'Bengaluru — 560 038, Karnataka\n'
        'kensora@kalakruthi.com\n'
        '+91 98765 43210'
    )
    rr.font.size      = Pt(8.5)
    rr.font.color.rgb = RGBColor(0xCC, 0xC5, 0xB8)
    rr.font.name      = 'Calibri'

    doc.add_paragraph()   # spacing after header

    # ════════════════════════════════════════════════════════════════════════
    #  DOCUMENT TITLE
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('Client Welcome & Project Proposal')
    r.font.size      = Pt(22)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = False
    r.font.name      = 'Palatino Linotype'

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run('Interior Design & 3D Visualisation Package')
    r2.font.size      = Pt(11)
    r2.font.color.rgb = GOLD
    r2.font.italic    = True
    r2.font.name      = 'Calibri'

    add_gold_rule(doc)

    # ════════════════════════════════════════════════════════════════════════
    #  META TABLE (Doc no | Date | Project ref | Version)
    # ════════════════════════════════════════════════════════════════════════
    mt = doc.add_table(rows=1, cols=4)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_border(mt, color_hex='E0D8CC', size='4')

    meta_fields = [
        ('Document Ref.',  'KEN-DOC-[XXXX]'),
        ('Prepared For',   '[Client Name]'),
        ('Date',           '[DD Month YYYY]'),
        ('Project Code',   'PRJ-[XXXX]-BLR'),
    ]
    for i, (label, val) in enumerate(meta_fields):
        c = mt.cell(0, i)
        set_cell_bg(c, 'F5F0E8')
        set_cell_margins(c, top=100, bottom=100, left=140, right=140)
        cp = c.paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rl = cp.add_run(label.upper() + '\n')
        rl.font.size      = Pt(7)
        rl.font.color.rgb = GOLD
        rl.font.bold      = True
        rl.font.name      = 'Calibri'
        rv = cp.add_run(val)
        rv.font.size      = Pt(10)
        rv.font.color.rgb = CHARCOAL
        rv.font.bold      = True
        rv.font.name      = 'Calibri'

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 1 — WELCOME LETTER
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('01  ·  A Warm Welcome from Kensora')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    body_para(doc,
        'Dear [Client Name],',
        bold=True, color=CHARCOAL, size_pt=11, space_after=8)

    body_para(doc,
        'On behalf of the entire Kensora by Kalakruthi team, we extend our warmest welcome '
        'and sincere gratitude for placing your trust in us to transform your home into a '
        'living work of art. This document marks the beginning of what we believe will be '
        'an extraordinary journey — one built on collaboration, craftsmanship, and a shared '
        'pursuit of beauty.',
        space_after=8)

    body_para(doc,
        'Kensora is the ultra-premium interior design division of Kalakruthi — a name '
        'synonymous with excellence in construction across South India. Founded on the belief '
        'that truly great interiors begin with understanding how you live, we bring together '
        'world-class materials, master artisans, and cutting-edge 3D visualisation tools to '
        'deliver spaces that are uniquely and unmistakably yours.',
        space_after=8)

    body_para(doc,
        'This proposal document outlines your project details, the requirements you have '
        'shared with us, and our preliminary design recommendations for each room of your '
        'residence. It also includes interactive 3D model references and Augmented Reality '
        '(AR) previews so you may experience your future spaces before a single item is '
        'moved.',
        space_after=8)

    body_para(doc,
        'We look forward to crafting something exceptional together.',
        space_after=4, italic=True, color=GOLD)

    body_para(doc, 'Warm regards,', color=CHARCOAL, space_after=2)
    body_para(doc, 'Ananya Sharma  ·  Design Director', color=CHARCOAL, bold=True, space_after=2)
    body_para(doc, 'Kensora by Kalakruthi', color=TEAL_DARK, space_after=2)
    body_para(doc, 'kensora@kalakruthi.com  ·  +91 98765 43210', color=TEXT_MID, size_pt=9.5)

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 2 — PROJECT DETAILS
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('02  ·  Project Details')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    # Project details table
    pd_tbl = doc.add_table(rows=8, cols=2)
    pd_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    add_border(pd_tbl, color_hex='B59A6A', size='6')

    pd_rows = [
        ('Project Name',         '[Project / Property Name]'),
        ('Client Name',          '[Full Name of Client]'),
        ('Property Address',     '[Full Property Address]'),
        ('Property Type',        '[Apartment / Villa / Penthouse / Duplex]'),
        ('Total Area',           '[XXX] sq.ft  ·  [X] BHK'),
        ('Project Scope',        '[Full Home / Selected Rooms]'),
        ('Design Style',         '[Contemporary / Classic / Transitional / Bespoke]'),
        ('Estimated Timeline',   '[X] months  (Start: [Month YYYY]  |  Handover: [Month YYYY])'),
    ]
    for i, (k, v) in enumerate(pd_rows):
        kc = pd_tbl.cell(i, 0)
        vc = pd_tbl.cell(i, 1)
        bg = 'F5F0E8' if i % 2 == 0 else 'FDFBF8'
        set_cell_bg(kc, bg); set_cell_bg(vc, bg)
        set_cell_margins(kc, top=90, bottom=90, left=140, right=80)
        set_cell_margins(vc, top=90, bottom=90, left=80, right=140)

        kp = kc.paragraphs[0]
        kr = kp.add_run(k)
        kr.font.size      = Pt(9.5)
        kr.font.color.rgb = GOLD
        kr.font.bold      = True
        kr.font.name      = 'Calibri'

        vp = vc.paragraphs[0]
        vr = vp.add_run(v)
        vr.font.size      = Pt(10)
        vr.font.color.rgb = CHARCOAL
        vr.font.name      = 'Calibri'

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 3 — OVERALL CLIENT REQUIREMENTS
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('03  ·  Overall Client Requirements')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    body_para(doc,
        'The following requirements have been captured during the initial consultation. '
        'Please review and confirm, or annotate any changes before the design phase begins.',
        space_after=10)

    overall_requirements = [
        'Style Preference: [Client described preferred aesthetic — e.g., Warm Contemporary with natural materials]',
        'Colour Palette: [Preferred tones — e.g., Warm Neutrals, Deep Teals, Gold accents]',
        'Material Preferences: [e.g., Marble, Solid Wood, Linen Fabrics, Brushed Brass fixtures]',
        'Lifestyle Needs: [e.g., WFH setup, children-friendly spaces, entertainment zone, home gym]',
        'Smart Home Integration: [Yes / No]  — [Specify systems if applicable]',
        'Sustainability Focus: [Level of preference for eco-friendly and sustainably sourced materials]',
        'Budget Range: ₹ [X,XX,XXX] – ₹ [X,XX,XXX] (indicative)',
        'Key Priority Rooms: [Rooms the client emphasises most]',
        'Special Requirements: [Any accessibility needs, heirloom furniture to retain, art to integrate, etc.]',
    ]
    for req in overall_requirements:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after  = Pt(4)
        p.paragraph_format.space_before = Pt(0)
        r = p.add_run(req)
        r.font.size      = Pt(10)
        r.font.color.rgb = TEXT_MID
        r.font.name      = 'Calibri'

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 4 — ROOM-BY-ROOM DESIGN PROPOSALS
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('04  ·  Room-by-Room Design Proposals')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    body_para(doc,
        'Each room below includes your stated requirements, a 3D model visualisation placeholder '
        '(for the rendered digital model your designer will attach), an Augmented Reality (AR) '
        'preview placeholder, and Kensora\'s curated design recommendations.',
        space_after=12)

    # ── Rooms Data ────────────────────────────────────────────────────────────
    rooms = [
        {
            'name':  'Living Room',
            'icon':  '🛋',
            'requirements': [
                '[Requirement 1 — e.g., Open-plan layout connecting to dining area]',
                '[Requirement 2 — e.g., Feature wall with fireplace or statement artwork]',
                '[Requirement 3 — e.g., Seating for 8–10 guests with a home bar unit]',
                '[Requirement 4 — e.g., Integrated entertainment unit with concealed wiring]',
            ],
            'usecase_3d': (
                'This 3D model presents the full-length view of the living area from the entryway, '
                'showcasing the spatial layout, furniture arrangement, lighting plan, and proposed '
                'material finishes. Use this view to evaluate proportions, traffic flow, and overall '
                'ambiance before confirming the design direction.'
            ),
            'usecase_ar': (
                'Scan the QR code or launch the AR link on your smartphone to walk through the '
                'living room virtually at true scale. Reposition furniture, swap material finishes '
                'in real time, and share the AR session with family members for collaborative '
                'decision making — all before any item is ordered.'
            ),
            'recommendations': [
                'Material: Vein-matched Statuario marble flooring with underfloor heating for year-round luxury comfort.',
                'Furniture: Bespoke low-profile sectional sofa in Belgian linen — custom upholstered by Kensora artisans.',
                'Lighting: Layered lighting scheme — recessed wash, sculptural pendant over conversation zone, concealed cove lighting.',
                'Feature Wall: Micro-textured Venetian plaster panel in dusty rose or deep slate, flanked by custom fluted panelling.',
                'Storage: Full-height fluted walnut joinery unit with integrated display niches and concealed AV cabinet.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Master Bedroom',
            'icon':  '🛏',
            'requirements': [
                '[Requirement 1 — e.g., King-size bed with upholstered headboard wall]',
                '[Requirement 2 — e.g., Walk-in wardrobe connecting to and integrated with the en-suite]',
                '[Requirement 3 — e.g., Reading nook or private seating area by window]',
                '[Requirement 4 — e.g., Smart lighting and blackout drapery]',
            ],
            'usecase_3d': (
                'The 3D model renders the master bedroom from the door threshold, capturing the '
                'bed wall composition, wardrobe facade, window treatment, and ceiling detail. '
                'Review the scale of furniture against room dimensions and confirm material '
                'selections for the headboard, joinery, and flooring.'
            ),
            'usecase_ar': (
                'Use the AR model to superimpose the designed master bedroom into your actual '
                'existing space. Walk around the virtual bed and wardrobe to verify spatial '
                'clearances, evaluate how natural light interacts with the proposed finishes at '
                'different times of day, and confirm the palette before committing.'
            ),
            'recommendations': [
                'Headboard Wall: Full-height fluted oak panelling with integrated fabric panel in stone-washed linen.',
                'Flooring: Wide-plank French oak engineered timber with herringbone border detail.',
                'Wardrobe: Floor-to-ceiling bespoke wardrobe in lacquered matte ivory with brushed gold hardware.',
                'Ceiling: Coffered plaster ceiling with concealed warm LED strip — dimmable to 2700K.',
                'Drapery: Motorised blackout lining with sheer outer panel in ivory silk organza.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Kitchen & Dining',
            'icon':  '🍽',
            'requirements': [
                '[Requirement 1 — e.g., Open kitchen with island and breakfast bar]',
                '[Requirement 2 — e.g., High-end appliances (specify brands if preferred)]',
                '[Requirement 3 — e.g., Separate dry and wet kitchen zones]',
                '[Requirement 4 — e.g., Dining table seating for 8 persons]',
            ],
            'usecase_3d': (
                'The kitchen 3D model provides a bird\'s-eye and eye-level view of the full '
                'kitchen-dining zone, showing cabinet layout, worktop material, appliance '
                'placement, island dimensions, and the dining setting. Use this model to '
                'finalise the functional workflow (cook\'s triangle) and storage provisions.'
            ),
            'usecase_ar': (
                'The AR kitchen model allows you to physically walk around the proposed island, '
                'open virtual cabinet doors, and test the ergonomic reach to overhead storage. '
                'Ideal for clients who wish to validate the functional layout before fabrication '
                'begins — saving costly revisions during installation.'
            ),
            'recommendations': [
                'Worktop: Calacatta Gold quartz with waterfall edge detail on kitchen island.',
                'Cabinetry: Two-tone scheme — upper cabinets in high-gloss lacquer, lower in hand-painted shaker with brass cup pulls.',
                'Backsplash: Handcrafted zellige tile in sage or off-white — adds artisanal texture to the cooking zone.',
                'Appliances: Miele or Gaggenau integrated suite — fully panel-matched for a seamless aesthetic.',
                'Dining: Live-edge solid walnut dining table with custom Kensora upholstered chairs in bouclé.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Home Office / Study',
            'icon':  '📐',
            'requirements': [
                '[Requirement 1 — e.g., Built-in desk with dual monitor setup]',
                '[Requirement 2 — e.g., Acoustically treated walls for video calls]',
                '[Requirement 3 — e.g., Dedicated bookshelf and display for awards / art]',
                '[Requirement 4 — e.g., Ergonomic task lighting and natural light access]',
            ],
            'usecase_3d': (
                'The home office 3D model details the desk configuration, built-in shelving '
                'system, acoustic treatment placement, and lighting scheme. Evaluate the visual '
                'weight of joinery against the room\'s proportions and confirm the balance between '
                'functional work surfaces and aesthetic display areas.'
            ),
            'usecase_ar': (
                'Place the virtual home office into your actual room using AR to test desk '
                'positioning relative to windows (to avoid screen glare), confirm the ergonomic '
                'chair reach to storage, and experience the acoustic panel textures at human scale '
                'before installation.'
            ),
            'recommendations': [
                'Desk: Custom floating desk in matte black lacquer with integrated cable tray — 240 cm wide for executive presence.',
                'Shelving: Full-height open shelving in smoked oak with integrated task lighting and display ledges.',
                'Acoustic Panels: Fabric-wrapped acoustic art panels in premium linen — dual function: beauty and sound control.',
                'Chair: Provide allowance for client\'s preferred ergonomic task chair; Kensora to specify complementary visitor seating.',
                'Lighting: Articulated wall sconce for task reading, concealed LED troffers for ambient, pendant for atmosphere.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Children\'s Room',
            'icon':  '🎨',
            'requirements': [
                '[Requirement 1 — e.g., Twin beds or bunk bed configuration]',
                '[Requirement 2 — e.g., Study zone with individual desks]',
                '[Requirement 3 — e.g., Play area with safe, durable flooring]',
                '[Requirement 4 — e.g., Ample storage for toys, books, sportswear]',
            ],
            'usecase_3d': (
                'The children\'s room 3D model visualises the safe and stimulating layout — '
                'showing bed arrangement, study corner, play zone, and storage wall. Parents '
                'can evaluate traffic clearances between furniture, confirm the playful yet '
                'sophisticated palette, and visualise growth adaptability (the design should '
                'age with the child).'
            ),
            'usecase_ar': (
                'AR allows parents and children alike to walk through the room together virtually, '
                'sparking engagement and excitement about their new space. Test that window '
                'positioning supports natural light for the study zone, and that the play floor '
                'area meeting minimum safety clearances.'
            ),
            'recommendations': [
                'Flooring: Luxury cork flooring — warm, soft, acoustically dampening, and impact-resilient for active children.',
                'Furniture: Modular bed system with under-bed storage drawers; convertible desk-to-display unit as child grows.',
                'Palette: Calm, sophisticated tones — dusty sage, warm terracotta, natural oak — avoids visually overwhelming primaries.',
                'Storage: Customised open shelving at child\'s reachable height with labelled fabric baskets.',
                'Safety: All furniture with rounded corners; wall-mounted items with concealed fixings rated to 5× working load.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Master En-Suite Bathroom',
            'icon':  '🛁',
            'requirements': [
                '[Requirement 1 — e.g., Freestanding bathtub and walk-in rainfall shower]',
                '[Requirement 2 — e.g., Double vanity with hotel-style illuminated mirrors]',
                '[Requirement 3 — e.g., Heated floors and towel rails]',
                '[Requirement 4 — e.g., Separate WC compartment]',
            ],
            'usecase_3d': (
                'The bathroom 3D model presents the luxurious en-suite layout — bathtub '
                'positioning, shower enclosure detailing, vanity elevation, and tile pattern '
                'arrangement. Verify that the spatial experience aligns with your expectations '
                'of a spa-like sanctuary before any wet works commence.'
            ),
            'usecase_ar': (
                'Walk into your bathroom virtually using the AR model. Confirm the visual '
                'weight of the freestanding bath, check that the shower enclosure provides '
                'adequate turning radius, and compare tile options in the true space and '
                'lighting of your actual bathroom.'
            ),
            'recommendations': [
                'Bath: Freestanding sculptural bathtub in matte stone or Corian — positioned to celebrate window view.',
                'Shower: Full-height fluted glass enclosure with frameless fittings; overhead rain head + body jets.',
                'Vanity: Floating double vanity in honed Nero Marquina marble top with undermount sinks; wall-hung taps.',
                'Tiles: Large-format porcelain slabs in bookmatched vein pattern — grout-minimal for easy maintenance.',
                'Accessories: Brushed unlacquered brass fixtures throughout for warmth and patina over time.',
                '[Add additional Kensora recommendation here]',
            ],
        },
        {
            'name':  'Entryway / Foyer',
            'icon':  '🚪',
            'requirements': [
                '[Requirement 1 — e.g., Statement entry console and mirror]',
                '[Requirement 2 — e.g., Shoe storage concealed within joinery]',
                '[Requirement 3 — e.g., Art display or sculptural feature]',
                '[Requirement 4 — e.g., Smart entry with intercom / access control]',
            ],
            'usecase_3d': (
                'The foyer 3D model captures the all-important first impression of the home — '
                'front door treatment, console styling, overhead pendant, feature wall, and '
                'flooring inlay transition from entry to living. This view sets the tone for '
                'the entire residence and should be reviewed with particular attention to '
                'material contrast and lighting drama.'
            ),
            'usecase_ar': (
                'Use the AR foyer model to experience the entry sequence — open the virtual '
                'door and see the space unfold. Test how the pendant light scale reads against '
                'ceiling height, confirm the console proportions against the wall width, and '
                'evaluate the flooring inlay geometry at true scale.'
            ),
            'recommendations': [
                'Flooring: Decorative marble inlay or brass-inlaid geometric pattern to signal arrival.',
                'Console: Bespoke lacquered console with hand-cast brass legs and integrated hidden storage.',
                'Pendant: Statement hand-blown glass chandelier or sculptural pendant — the foyer\'s crown jewel.',
                'Mirror: Oversized arched mirror in aged brass frame — amplifies space and light.',
                'Art: Commission a Kensora-curated piece from our partner gallery network for the feature wall.',
                '[Add additional Kensora recommendation here]',
            ],
        },
    ]

    for room in rooms:
        add_room_section(
            doc,
            room_name      = room['name'],
            room_icon      = room['icon'],
            requirements   = room['requirements'],
            usecase_3d     = room['usecase_3d'],
            usecase_ar     = room['usecase_ar'],
            recommendations= room['recommendations'],
        )
        # Page break between rooms (not after the last)
        if room != rooms[-1]:
            doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 5 — NEXT STEPS
    # ════════════════════════════════════════════════════════════════════════
    doc.add_page_break()

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('05  ·  Next Steps')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    steps = [
        ('Step 1 — Review & Annotate',
         'Please review this document carefully. Add your comments, corrections, or approvals '
         'directly onto this file or share feedback via email / WhatsApp to your assigned designer.'),
        ('Step 2 — Design Concept Presentation',
         'Your Kensora designer will schedule a concept presentation session where mood boards, '
         'material swatches, and refined 3D views are presented for all key rooms.'),
        ('Step 3 — Material Approval & FF&E Selection',
         'Selected materials, furniture, fixtures, and equipment will be presented for approval. '
         'Kensora will provide sample boards for all primary surfaces.'),
        ('Step 4 — Working Drawings & Estimates',
         'Detailed working drawings, shop drawings, and a phased cost estimate will be shared '
         'for your sign-off before any fabrication or procurement commences.'),
        ('Step 5 — Execution & Site Supervision',
         'Kensora\'s project management team oversees all on-site work, vendor coordination, '
         'quality inspections, and delivery milestones to ensure seamless execution.'),
        ('Step 6 — Handover & Styling',
         'The final handover includes a complete styling session, a snag-list walkthrough, '
         'and a dedicated after-care briefing for all installed systems and materials.'),
    ]

    for title, desc in steps:
        label_para(doc, title)
        body_para(doc, desc, space_after=6)

    # ════════════════════════════════════════════════════════════════════════
    #  SECTION 6 — TERMS & CONFIDENTIALITY
    # ════════════════════════════════════════════════════════════════════════
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run('06  ·  Confidentiality & Document Terms')
    r.font.size      = Pt(16)
    r.font.color.rgb = TEAL_DARK
    r.font.bold      = True
    r.font.name      = 'Palatino Linotype'

    add_gold_rule(doc)

    body_para(doc,
        'This document is prepared exclusively for the client named herein and is strictly '
        'confidential. The design concepts, recommendations, 3D visualisations, and AR models '
        'contained within remain the intellectual property of Kensora by Kalakruthi until a '
        'formal design agreement is executed and the retainer payment is received.',
        space_after=8)

    body_para(doc,
        'Reproduction, distribution, or use of any part of this document by a third party '
        'without prior written consent from Kensora by Kalakruthi is strictly prohibited.',
        space_after=8)

    body_para(doc,
        'All disputes are subject to the exclusive jurisdiction of the courts of Bengaluru, Karnataka.',
        space_after=16)

    # Signature block
    sig_tbl = doc.add_table(rows=1, cols=2)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ci, (name, role) in enumerate([
        ('Client Signature', '[Client Name]'),
        ('Kensora Authorisation', 'Design Director · Kensora by Kalakruthi'),
    ]):
        sc = sig_tbl.cell(0, ci)
        set_cell_bg(sc, 'FDFBF8')
        set_cell_margins(sc, top=160, bottom=160, left=200, right=200)
        sp1 = sc.paragraphs[0]
        sp1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Signature line
        sr1 = sp1.add_run('_' * 36)
        sr1.font.size      = Pt(10)
        sr1.font.color.rgb = RGBColor(0xCC, 0xC5, 0xB8)
        sr1.font.name      = 'Calibri'
        sp2 = sc.add_paragraph()
        sp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr2 = sp2.add_run(name)
        sr2.font.size      = Pt(8.5)
        sr2.font.color.rgb = GOLD
        sr2.font.bold      = True
        sr2.font.name      = 'Calibri'
        sp3 = sc.add_paragraph()
        sp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr3 = sp3.add_run(role)
        sr3.font.size      = Pt(9)
        sr3.font.color.rgb = CHARCOAL
        sr3.font.name      = 'Calibri'

    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════════════
    #  LETTERHEAD FOOTER
    # ════════════════════════════════════════════════════════════════════════
    ft = doc.add_table(rows=1, cols=1)
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    fc = ft.cell(0, 0)
    set_cell_bg(fc, '0D3D3A')
    set_cell_margins(fc, top=150, bottom=150, left=300, right=300)
    fp = fc.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        'Kensora by Kalakruthi  ·  Design Studio, Indiranagar, Bengaluru — 560 038\n'
        'kensora@kalakruthi.com  ·  +91 98765 43210  ·  www.kensora.in\n'
        '◈  Ultra-Premium Interiors  ◈  India Design Award 2024  ◈  Luxury Interior Guild'
    )
    fr.font.size      = Pt(8)
    fr.font.color.rgb = RGBColor(0xBB, 0xB5, 0xAC)
    fr.font.name      = 'Calibri'

    # ────────────────────────────────────────────────────────────────────────
    out_path = r'd:\GIT\KENSORA\Kensora_Client_Welcome_Template.docx'
    doc.save(out_path)
    print(f'[OK]  Document saved -> {out_path}')


if __name__ == '__main__':
    build_document()
